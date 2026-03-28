from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
import hashlib


@dataclass
class Document:
    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: Optional[str] = None

    def __post_init__(self):
        if self.doc_id is None:
            self.doc_id = hashlib.md5(self.content.encode()).hexdigest()[:12]


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".md", ".markdown", ".txt",
        ".py", ".js", ".ts", ".java", ".cpp", ".c", ".h",
        ".go", ".rs", ".rb", ".php", ".swift", ".kt",
        ".json", ".yaml", ".yml", ".xml", ".html", ".css"
    }

    CODE_EXTENSIONS = {
        ".py", ".js", ".ts", ".java", ".cpp", ".c", ".h",
        ".go", ".rs", ".rb", ".php", ".swift", ".kt",
        ".json", ".yaml", ".yml", ".xml", ".html", ".css"
    }

    def load(self, file_path: str | Path) -> Document:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        metadata = {
            "source": str(path.absolute()),
            "filename": path.name,
            "extension": ext,
        }

        if ext == ".pdf":
            content = self._load_pdf(path)
            metadata["type"] = "pdf"
        elif ext == ".docx":
            content = self._load_docx(path)
            metadata["type"] = "docx"
        elif ext in {".md", ".markdown"}:
            content = self._load_text(path)
            metadata["type"] = "markdown"
        elif ext in self.CODE_EXTENSIONS:
            content = self._load_text(path)
            metadata["type"] = "code"
            metadata["language"] = self._get_language(ext)
        else:
            content = self._load_text(path)
            metadata["type"] = "text"

        return Document(content=content, metadata=metadata)

    def load_directory(self, dir_path: str | Path, recursive: bool = True) -> list[Document]:
        path = Path(dir_path)

        if not path.is_dir():
            raise NotADirectoryError(f"Not a directory: {path}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load(file_path)
                    documents.append(doc)
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")

        return documents

    def _load_pdf(self, path: Path) -> str:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF support. Install with: pip install pypdf2")

        reader = PdfReader(path)
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f"[Page {page_num}]\n{text}")

        return "\n\n".join(text_parts)

    def _load_docx(self, path: Path) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")

        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_text(self, path: Path) -> str:
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue

        raise ValueError(f"Could not decode file: {path}")

    def _get_language(self, ext: str) -> str:
        lang_map = {
            ".py": "python",
            ".js": "javascript",
            ".ts": "typescript",
            ".java": "java",
            ".cpp": "cpp",
            ".c": "c",
            ".h": "c",
            ".go": "go",
            ".rs": "rust",
            ".rb": "ruby",
            ".php": "php",
            ".swift": "swift",
            ".kt": "kotlin",
            ".json": "json",
            ".yaml": "yaml",
            ".yml": "yaml",
            ".xml": "xml",
            ".html": "html",
            ".css": "css",
        }
        return lang_map.get(ext, "unknown")
